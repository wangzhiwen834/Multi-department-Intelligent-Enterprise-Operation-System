# step1: front matter (title, authors, abstract, keywords)
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
style.paragraph_format.line_spacing = Pt(20)
for name, size in [('Heading 1', Pt(16)), ('Heading 2', Pt(14)), ('Heading 3', Pt(12))]:
    h = doc.styles[name]
    h.font.size = size; h.font.bold = True; h.font.color.rgb = RGBColor(0,0,0)
    h._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h.paragraph_format.line_spacing = Pt(20)

def add_text(p, text, bold=False, italic=False, font_cn='宋体', font_en='Times New Roman', size=Pt(12)):
    run = p.add_run(text)
    run.font.name = font_en; run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
    run.font.size = size; run.bold = bold; run.italic = italic

def add_para(text, style_name='Normal', align=None, first_indent=Cm(0.74)):
    p = doc.add_paragraph(style=style_name)
    if align: p.alignment = align
    p.paragraph_format.first_line_indent = first_indent if style_name == 'Normal' else Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    import re
    parts = re.split(r'(\*\*.+?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'): add_text(p, part[2:-2], bold=True)
        elif part.startswith('`') and part.endswith('`'): add_text(p, part[1:-1], font_cn='Courier New', font_en='Courier New')
        else: add_text(p, part)
    return p

def add_heading(text, level):
    names = {0: 'Heading 1', 1: 'Heading 2', 2: 'Heading 3'}
    p = doc.add_paragraph(style=names[level])
    add_text(p, text, font_cn='黑体'); p.paragraph_format.first_line_indent = Cm(0)

# 基金项目
fp = doc.add_paragraph(); add_text(fp, '基金项目：', bold=True); add_text(fp, '[请填写，如无则删除本段]'); fp.paragraph_format.first_line_indent = Cm(0)
# 作者简介
fp = doc.add_paragraph(); add_text(fp, '作者简介：', bold=True); add_text(fp, '[第一作者姓名]（[出生年]-），[性别]，[民族]，[籍贯（市）]，[职称]，[学历]，研究方向：[研究方向]。'); fp.paragraph_format.first_line_indent = Cm(0)
doc.add_paragraph()
# 标题
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(tp, '多部门智能经营系统的设计与实现', font_cn='黑体', font_en='Times New Roman', size=Pt(16), bold=True)
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(tp, '[作者姓名1]，[作者姓名2]，[作者姓名1,*]', size=Pt(10.5))
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(tp, '（[作者单位全称]  [省市]  [邮编]）', size=Pt(10.5))
doc.add_paragraph()
# 摘要
add_para('摘  要：中小企业经营数据长期依赖散乱Excel与纸质笔记，存在跨表公式脆弱与重复录入的突出问题。针对该问题，设计并实现了一个基于Univer在线表格与豆包大模型的多业态智能经营系统。系统采用"录入与展示分离"架构，以PostgreSQL为唯一事实源、Univer为录入面；以AI语义抽取作为唯一入库路径，以大模型按表头文字语义对齐字段替代位置式同步，实现抗布局变化、跨业务可复用、近确定性可复现的经营数据入库；并以工作表级悲观锁替代操作变换实现轻量协同录入。系统已部署上线并交付真实门店使用，以足浴门店真实经营数据验证了抽取管线的有效性，遗留Excel的公式错误值被校验机制正确拦截。研究表明大语言模型可作为结构化数据入库的语义对齐层，为中小企业经营数字化提供了可复用的设计模式。', first_indent=Cm(0))
# 关键词
kp = doc.add_paragraph(); add_text(kp, '关键词：', bold=True); add_text(kp, '中小企业；经营数据；大语言模型；信息抽取；在线电子表格；悲观锁'); kp.paragraph_format.first_line_indent = Cm(0)
doc.add_paragraph()

doc.save('docs/论文-中国商论版.docx')
print('step1 done')
