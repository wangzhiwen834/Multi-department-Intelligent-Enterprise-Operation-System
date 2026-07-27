# 论文新版 -> 汉斯《计算机科学与应用》模板排版 (基于 to-hans.py,适配 docs/论文-正文-新版.md)
# 用法: python docs/to-hans-new.py   (从仓库根目录运行)
# 改动 vs to-hans.py:读新版草稿;前置从 md 解析;跳过前置至 ## 1;处理 **表N** + | 表格;
#   FIGS 按出现顺序重编号;剥离 \tag 防双编号;文末补致谢/基金项目/参考文献占位。
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import re, os, subprocess, tempfile, zipfile, copy, shutil

# ============ pandoc LaTeX -> 原生 OMML ============
def _find_pandoc():
    p = shutil.which('pandoc')
    if p: return p
    for cand in (r"C:\Users\11422\AppData\Local\Pandoc\pandoc.exe", r"C:\Program Files\Pandoc\pandoc.exe"):
        if os.path.exists(cand): return cand
    return None

PANDOC = _find_pandoc()
_M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

def _latex_to_omath(latex):
    if not PANDOC: return None
    with tempfile.TemporaryDirectory() as td:
        mdp = os.path.join(td, 'eq.md'); docxp = os.path.join(td, 'eq.docx')
        with open(mdp, 'w', encoding='utf-8') as f: f.write('$' + latex + '$')
        r = subprocess.run([PANDOC, mdp, '-f', 'markdown', '-t', 'docx', '-o', docxp], capture_output=True)
        if r.returncode != 0 or not os.path.exists(docxp): return None
        with zipfile.ZipFile(docxp) as z: xml = z.read('word/document.xml')
    tree = etree.fromstring(xml)
    return tree.find('.//{%s}oMath' % _M_NS)

_EQ_N = [0]
def add_equation(doc, latex):
    _EQ_N[0] += 1
    p = doc.add_paragraph(style='Normal')
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0); pf.left_indent = Cm(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(4); pf.space_after = Pt(4)
    pf.tab_stops.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.RIGHT)
    omath = _latex_to_omath(latex)
    if omath is not None:
        p._element.append(copy.deepcopy(omath))
    else:
        rb = p.add_run(latex); rb.font.name = 'Cambria Math'
        rb._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); rb.font.size = Pt(10.5)
    r2 = p.add_run(); r2.add_tab()
    r3 = p.add_run('(%d)' % _EQ_N[0]); r3.font.size = Pt(10.5)
    return p

# ============ 文档与样式 ============
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

def setfont(run, ascii_font, ea_font, size, bold=None, italic=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), ea_font)
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def para(text, ascii_font, ea_font, size, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent is not None: p.paragraph_format.first_line_indent = Cm(indent)
    r = p.add_run(text); setfont(r, ascii_font, ea_font, size, bold)
    return p

# ============ 从新版草稿解析前置(标题/摘要/关键词) ============
SRC = 'docs/论文-正文-新版.md'
with open(SRC, encoding='utf-8') as f:
    src = f.read()

def first_match(pattern, flags=re.M):
    m = re.search(pattern, src, flags)
    return m.group(1).strip() if m else ''

TITLE_CN = first_match(r'^# (.+)$')
TITLE_EN = first_match(r'^\*\*(.+?)\*\*\s*$')   # 整行为 **...** 的首行(英文标题)
ABSTRACT_CN = first_match(r'^\*\*摘要[:：]\*\*\s*(.+)$')
KEYWORDS_CN = first_match(r'^\*\*关键词[:：]\*\*\s*(.+)$')
ABSTRACT_EN = first_match(r'^\*\*Abstract[:：]\*\*\s*(.+)$')
KEYWORDS_EN = first_match(r'^\*\*Keywords[:：]\*\*\s*(.+)$')

# ============ 中文前置 ============
para(TITLE_CN, 'Times New Roman', '黑体', 22, bold=True)
para('作者姓名1，作者姓名2', 'Times New Roman', '楷体_GB2312', 11, bold=True)
para('（单位全称 院/系/科室，省 市 邮编）', 'Times New Roman', '宋体', 10)
para('收稿日期：******', 'Times New Roman', '宋体', 10)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
setfont(p.add_run('摘要：'), 'Times New Roman', '黑体', 12, bold=True)
setfont(p.add_run(ABSTRACT_CN), 'Times New Roman', '宋体', 10)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(12); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
setfont(p.add_run('关键词：'), 'Times New Roman', '黑体', 12, bold=True)
setfont(p.add_run(KEYWORDS_CN), 'Times New Roman', '宋体', 10)

# ============ 英文前置 ============
para(TITLE_EN, 'Times New Roman', '宋体', 22, bold=True, before=12)
para('Author Name1, Author Name2', 'Times New Roman', '宋体', 11, bold=True)
para('(Dept. name of organization, City, Postal Code, Country)', 'Times New Roman', '宋体', 10)
para('Received: ******', 'Times New Roman', '宋体', 10)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
setfont(p.add_run('Abstract: '), 'Times New Roman', '宋体', 12, bold=True)
setfont(p.add_run(ABSTRACT_EN), 'Times New Roman', '宋体', 10)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(12); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
setfont(p.add_run('Keywords: '), 'Times New Roman', '宋体', 12, bold=True)
setfont(p.add_run(KEYWORDS_EN), 'Times New Roman', '宋体', 10)

para('Copyright © 2026 by author(s) and Hans Publishers Inc.', 'Times New Roman', '宋体', 9, before=6)
para('This work is licensed under the Creative Commons Attribution International License (CC BY 4.0).', 'Times New Roman', '宋体', 9)
para('http://creativecommons.org/licenses/by/4.0/', 'Times New Roman', '宋体', 9)

# ============ 正文解析 ============
with open(SRC, encoding='utf-8') as f:
    lines = f.readlines()

def heading_text(text):
    m = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', text)
    return '%s. %s' % (m.group(1), m.group(2)) if m else text

def add_heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6 if level < 3 else 0)
    r = p.add_run(heading_text(text))
    sz = {1: 12, 2: 11, 3: 10.5}[level]
    setfont(r, 'Times New Roman', '黑体', sz, bold=True)
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    parts = re.split(r'(\*\*.+?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); setfont(r, 'Times New Roman', '宋体', 10.5, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); setfont(r, 'Courier New', '宋体', 9)
        else:
            r = p.add_run(part); setfont(r, 'Times New Roman', '宋体', 10.5)
    return p

def add_reference(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0); p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.hanging_indent = Cm(0.74); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text); setfont(r, 'Times New Roman', '宋体', 9)
    return p

def add_table_caption(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); setfont(r, 'Times New Roman', '黑体', 9, bold=True)

def add_md_table(table_lines):
    rows = []
    for ln in table_lines:
        cells = [c.strip() for c in ln.strip().strip('|').split('|')]
        if cells and all(re.match(r'^[-: ]*$', c) for c in cells):  # 分隔行
            continue
        rows.append(cells)
    if not rows: return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = tbl.cell(ri, ci); cell.text = ''
            pp = cell.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            txt = row[ci] if ci < len(row) else ''
            r = pp.add_run(txt); setfont(r, 'Times New Roman', '宋体', 9, bold=(ri == 0))
    para('', 'Times New Roman', '宋体', 9)

# 图号按出现顺序重编号:图1=架构/图2=ER/图3=抽取管线/图4=锁/图5=五用/图6=分派/图7=足浴评估
FIGS = {
    1: ('fig1-architecture.png', 'Figure 1. System Overall Architecture', '图1. 系统总体架构图'),
    2: ('fig3-er-diagram.png', 'Figure 2. Data Model ER Diagram', '图2. 数据模型ER图'),
    3: ('fig4-extraction-pipeline.png', 'Figure 3. AI Extraction Pipeline Flow', '图3. AI抽取管线流程图'),
    4: ('fig5-lock-state-machine.png', 'Figure 4. Pessimistic Lock State Machine', '图4. 悲观锁状态转换图'),
    5: ('fig6-template-five-uses.png', 'Figure 5. Template Descriptor: Five Reuses', '图5. 模板描述符“一份五用”'),
    6: ('fig7-business-dispatch.png', 'Figure 6. Multi-Business Dispatch', '图6. 多业务分派示意图'),
    7: ('fig2-footbath-eval.png', 'Figure 7. Real Business Data of Footbath Store (Jul 2026)', '图7. 足浴门店2026-07真实经营数据'),
}
_inserted = set()

def insert_figure(n):
    if n in _inserted or n not in FIGS: return
    _inserted.add(n)
    fname, cap_en, cap_cn = FIGS[n]
    path = 'docs/%s' % fname
    if not os.path.exists(path): return
    para('', 'Times New Roman', '宋体', 10)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.add_run().add_picture(path, width=Inches(5.0))
    para(cap_en, 'Times New Roman', '宋体', 9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=3)
    para(cap_cn, 'Times New Roman', '黑体', 9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para('', 'Times New Roman', '宋体', 10)

i = 0
started = False
saw_refs = False
while i < len(lines):
    line = lines[i].rstrip('\n'); i += 1
    if not line.strip(): continue
    if not started:
        if line.startswith('## '): started = True
        else: continue
    if line.startswith('# ') or line.startswith('> ') or line.startswith('---'): continue
    if line.startswith('<!--'): continue
    if line.startswith('**表'):  # 表标题
        add_table_caption(line.strip('*').strip()); continue
    if line.startswith('|'):  # markdown 表格
        tbl = [line]
        while i < len(lines) and lines[i].rstrip('\n').startswith('|'):
            tbl.append(lines[i].rstrip('\n')); i += 1
        add_md_table(tbl); continue
    m_eq = re.match(r'^\$\$(.+)\$\$$', line)
    if m_eq:
        latex = re.sub(r'\\tag\{[^}]*\}', '', m_eq.group(1)).strip()  # 剥离 \tag,由 add_equation 统一编号
        add_equation(doc, latex); continue
    if line.startswith('#### '): add_heading(line.replace('#### ', ''), 3); continue
    if line.startswith('### '): add_heading(line.replace('### ', ''), 2); continue
    if line.startswith('## '):
        htext = line.replace('## ', '')
        if htext.strip() == '参考文献' and not saw_refs:
            saw_refs = True
            add_heading('致谢', 1)
            para('（致谢信息。）', 'Times New Roman', '宋体', 10.5, indent=0.74)
            add_heading('基金项目', 1)
            para('（若为基金资助项目,请注明基金名称与编号。）', 'Times New Roman', '宋体', 10.5, indent=0.74)
        add_heading(htext, 1); continue
    if re.match(r'^\[\d+\]', line): add_reference(line); continue
    # 漏过前置行兜底跳过
    if line.startswith('**摘要') or line.startswith('**关键词') or line.startswith('**Abstract') or line.startswith('**Keywords'): continue
    if line.startswith('**') and line.endswith('**') and len(line) > 2: continue
    add_body(line)
    for mfig in re.finditer(r'图\s*(\d+)', line):
        insert_figure(int(mfig.group(1)))

# 未提及的图文末补插
for _n in sorted(FIGS):
    insert_figure(_n)

# 文末补致谢/基金项目/参考文献占位(若正文未含参考文献节)
if not saw_refs:
    add_heading('致谢', 1)
    para('（致谢信息。）', 'Times New Roman', '宋体', 10.5, indent=0.74)
    add_heading('基金项目', 1)
    para('（若为基金资助项目,请注明基金名称与编号。）', 'Times New Roman', '宋体', 10.5, indent=0.74)
    add_heading('参考文献', 1)
    para('[1] （待回填：GB/T 7714 著录，共 13 条，对应正文 [待补文献 1–13]。）', 'Times New Roman', '宋体', 9, indent=0.74)

# ============ 页面 ============
section = doc.sections[0]
section.page_width = Cm(21.0); section.page_height = Cm(28.5)
section.top_margin = Cm(3.0); section.bottom_margin = Cm(3.0)
section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
section.header_distance = Cm(1.5); section.footer_distance = Cm(1.5)

out = 'docs/论文-计算机科学与应用版-新版.docx'
doc.save(out)
print('Saved:', out)
print('TITLE_CN:', TITLE_CN)
print('TITLE_EN:', TITLE_EN)
print('ABSTRACT_CN len:', len(ABSTRACT_CN), '| ABSTRACT_EN len:', len(ABSTRACT_EN))
print('equations:', _EQ_N[0], '| figures inserted:', sorted(_inserted))
