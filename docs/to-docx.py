# 论文 Markdown -> DOCX 转换脚本 (python-docx)
# 用法: python docs/to-docx.py
# 前提: pip install python-docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import re
import os
import subprocess
import tempfile
import zipfile
import copy
import shutil

doc = Document()

# ---------- 样式设置 ----------
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
style.paragraph_format.line_spacing = Pt(20)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# 标题样式
for lvl, name, size, bold in [
    (0, 'Heading 1', Pt(16), True),
    (1, 'Heading 2', Pt(14), True),
    (2, 'Heading 3', Pt(12), True),
]:
    h = doc.styles[name]
    h.font.name = '黑体'
    h.font.size = size
    h.font.bold = bold
    h.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    h._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h.paragraph_format.line_spacing = Pt(20)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.space_before = Pt(6)

# ---------- 辅助函数 ----------
def add_para(text, style_name='Normal', bold=False, italic=False):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.first_line_indent = Cm(0 if style_name != 'Normal' else 0.74)
    return p

def add_heading_text(text, level):
    # level 0=H1, 1=H2, 2=H3
    name = f'Heading {level+1}'
    p = doc.add_paragraph(style=name)
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.first_line_indent = Cm(0)
    return p

# ---------- 显示公式 ($$...$$, LaTeX -> 原生 OMML via pandoc) ----------
def _find_pandoc():
    """优先 PATH,回退到 winget 默认安装位置。"""
    p = shutil.which('pandoc')
    if p:
        return p
    for cand in (r"C:\Users\11422\AppData\Local\Pandoc\pandoc.exe",
                 r"C:\Program Files\Pandoc\pandoc.exe"):
        if os.path.exists(cand):
            return cand
    return None

PANDOC = _find_pandoc()
_M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

def _latex_to_omath(latex):
    """LaTeX 数学串 -> <m:oMath> lxml 元素(失败返回 None)。"""
    if not PANDOC:
        return None
    md = '$' + latex + '$'   # 行内数学 -> 单个 <m:oMath>
    with tempfile.TemporaryDirectory() as td:
        mdp = os.path.join(td, 'eq.md'); docxp = os.path.join(td, 'eq.docx')
        with open(mdp, 'w', encoding='utf-8') as f:
            f.write(md)
        r = subprocess.run([PANDOC, mdp, '-f', 'markdown', '-t', 'docx', '-o', docxp],
                           capture_output=True)
        if r.returncode != 0 or not os.path.exists(docxp):
            return None
        with zipfile.ZipFile(docxp) as z:
            xml = z.read('word/document.xml')
    tree = etree.fromstring(xml)
    return tree.find('.//{%s}oMath' % _M_NS)

_EQ_N = [0]
def add_equation(latex):
    _EQ_N[0] += 1
    n = _EQ_N[0]
    p = doc.add_paragraph(style='Normal')
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    # 居中公式 + 右对齐编号:文本宽度 14.66cm,中点 7.33cm
    pf.tab_stops.add_tab_stop(Cm(7.33), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Cm(14.66), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(); r1.add_tab()                 # 前导 tab -> 居中
    omath = _latex_to_omath(latex)
    if omath is not None:
        p._element.append(copy.deepcopy(omath))     # 原生公式对象(OMML)
    else:
        rb = p.add_run(latex)                       # 回退:纯文本
        rb.font.name = 'Cambria Math'
        rb._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        rb.font.size = Pt(10.5)
    r2 = p.add_run(); r2.add_tab()                 # 编号 tab -> 右对齐
    r3 = p.add_run('(%d)' % n)
    r3.font.size = Pt(10.5)
    return p

# ---------- 解析正文 ----------
with open('docs/论文-正文草稿.md', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')
    i += 1
    if not line.strip():
        continue
    if line.startswith('# ') and not line.startswith('##'):
        add_heading_text(line.replace('# ', ''), 0)
        continue
    if line.startswith('## '):
        add_heading_text(line.replace('## ', ''), 1)
        continue
    if line.startswith('### '):
        add_heading_text(line.replace('### ', ''), 2)
        continue
    if line.startswith('**摘要**'):
        p = add_para('摘要: ', bold=True)
        p.runs[0].bold = True
        rest = line.replace('**摘要**:', '').strip()
        if rest:
            r = p.add_run(rest)
            r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        continue
    if line.startswith('**Abstract**'):
        p = add_para('Abstract: ', bold=True)
        rest = line.replace('**Abstract**:', '').strip()
        if rest:
            r = p.add_run(rest)
            r.font.name = 'Times New Roman'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        continue
    if line.startswith('**Keywords**'):
        p = add_para('Keywords: ', bold=True)
        rest = line.replace('**Keywords**:', '').strip()
        if rest:
            r = p.add_run(rest)
            r.font.name = 'Times New Roman'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        continue
    if line.startswith('---'):
        continue
    if line.startswith('> '):
        continue
    if line.startswith('*') and line.endswith('*'):
        continue
    if line.startswith('[') and 'DOI:' in line:
        # 参考文献
        p = add_para(line, style_name='Normal')
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.hanging_indent = Cm(0.74)
        continue
    if line.startswith('图 '):
        # 图题
        p = add_para(line, style_name='Normal')
        p.paragraph_format.first_line_indent = Cm(0)
        continue

    # 显示公式 $$...$$
    m_eq = re.match(r'^\$\$(.+)\$\$$', line)
    if m_eq:
        add_equation(m_eq.group(1).strip())
        continue

    # 普通段落(合并行内加粗 **text**)
    text = line
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # 去掉 **, 后面手动加粗
    # 简单处理:整段作为普通文本,保留 `code` 为等宽
    p = add_para('', style_name='Normal')
    # 分段处理 **bold** 和 `code`
    parts = re.split(r'(\*\*.+?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Courier New'
        else:
            r = p.add_run(part)
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(10.5)

# ---------- 插入图片 ----------
figs = [
    ('fig1-architecture.png', '图 1  系统总体架构图'),
    ('fig2-footbath-eval.png', '图 2  足浴门店 2026-07 真实经营数据'),
    ('fig3-er-diagram.png', '图 3  数据模型 ER 图'),
    ('fig4-extraction-pipeline.png', '图 4  AI 抽取管线流程图'),
    ('fig5-lock-state-machine.png', '图 5  悲观锁状态转换图'),
    ('fig6-template-five-uses.png', '图 6  模板描述符"一份五用"'),
    ('fig7-business-dispatch.png', '图 7  多业务分派示意图'),
]

for fname, caption in figs:
    path = f'docs/{fname}'
    if os.path.exists(path):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(5.8))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.font.size = Pt(9)
        cr.font.name = '宋体'
        cr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        cr.italic = True

# ---------- 页面设置 ----------
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

out = 'docs/论文-正文.docx'
doc.save(out)
print(f'Saved: {out}')
