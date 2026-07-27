# 论文 -> 汉斯出版社《计算机科学与应用》模板排版 (python-docx + pandoc 原生 OMML)
# 用法: python docs/to-hans.py   (从仓库根目录运行)
# 依据: docs/计算机科学与应用模板.docx
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import re
import os
import subprocess
import tempfile
import zipfile
import copy
import shutil

# ============ pandoc LaTeX -> 原生 OMML ============
def _find_pandoc():
    p = shutil.which('pandoc')
    if p:
        return p
    for cand in (r"C:\Users\11422\AppData\Local\Pandoc\pandoc.exe",
                 r"C:\Program Files\Pandoc\pandoc.exe"):
        if os.path.exists(cand):
            return cand
    return None

PANDOC = _find_pandoc()
_M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

def _latex_to_omath(latex):
    if not PANDOC:
        return None
    md = '$' + latex + '$'
    with tempfile.TemporaryDirectory() as td:
        mdp = os.path.join(td, 'eq.md'); docxp = os.path.join(td, 'eq.docx')
        with open(mdp, 'w', encoding='utf-8') as f:
            f.write(md)
        r = subprocess.run([PANDOC, mdp, '-f', 'markdown', '-t', 'docx', '-o', docxp],
                           capture_output=True)
        if r.returncode != 0 or not os.path.exists(docxp):
            return None
        with zipfile.ZipFile(docxp) as z:
            xml = z.read('word/document.xml')
    tree = etree.fromstring(xml)
    return tree.find('.//{%s}oMath' % _M_NS)

_EQ_N = [0]
def add_equation(doc, latex):
    _EQ_N[0] += 1
    p = doc.add_paragraph(style='Normal')
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0); pf.left_indent = Cm(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(4); pf.space_after = Pt(4)
    pf.tab_stops.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.RIGHT)
    omath = _latex_to_omath(latex)
    if omath is not None:
        p._element.append(copy.deepcopy(omath))
    else:
        rb = p.add_run(latex); rb.font.name = 'Cambria Math'
        rb._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); rb.font.size = Pt(10.5)
    r2 = p.add_run(); r2.add_tab()
    r3 = p.add_run('(%d)' % _EQ_N[0]); r3.font.size = Pt(10.5)
    return p

# ============ 文档与样式 ============
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)


def setfont(run, ascii_font, ea_font, size, bold=None, italic=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), ea_font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def para(text, ascii_font, ea_font, size, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
         before=0, after=0, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    r = p.add_run(text)
    setfont(r, ascii_font, ea_font, size, bold)
    return p


# ============ 中文前置 ============
TITLE_CN = '多部门智能经营系统的设计与实现'
TITLE_EN = 'Design and Implementation of a Multi-department Intelligent Enterprise Operation System'
para(TITLE_CN, 'Times New Roman', '黑体', 22, bold=True)                      # 主标题
para('作者姓名1，作者姓名2', 'Times New Roman', '楷体_GB2312', 11, bold=True)  # 作者
para('（单位全称 院/系/科室，省 市 邮编）', 'Times New Roman', '宋体', 10)     # 单位
para('收稿日期：******', 'Times New Roman', '宋体', 10)                        # 收稿日期

# 中文摘要
ABSTRACT_CN = ('中小企业经营数据长期依赖散乱 Excel 与纸质笔记,存在跨表公式脆弱与重复录入的突出问题。'
    '为此设计并实现了一个基于 Univer 在线表格与豆包大模型的多业态智能经营系统。系统采用"录入与展示分离"架构,'
    '以 PostgreSQL 为唯一事实源、Univer 为录入面;以 AI 语义抽取作为唯一入库路径,以大模型按表头文字语义对齐字段替代位置式同步,'
    '实现抗布局变化、跨业务可复用、近确定性可复现的经营数据入库;并以工作表级悲观锁替代操作变换实现轻量协同录入。'
    '系统已部署上线并交付真实门店使用,以足浴门店真实经营数据验证了抽取管线的有效性,遗留 Excel 的公式错误值被校验机制正确拦截。'
    '研究表明大语言模型可作为结构化数据入库的语义对齐层,为中小企业经营数字化提供了可复用的设计模式。')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
setfont(p.add_run('摘要：'), 'Times New Roman', '黑体', 12, bold=True)
setfont(p.add_run(ABSTRACT_CN), 'Times New Roman', '宋体', 10)
# 中文关键词
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(12); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
setfont(p.add_run('关键词：'), 'Times New Roman', '黑体', 12, bold=True)
setfont(p.add_run('中小企业；经营数据；大语言模型；信息抽取；在线表格；悲观锁'), 'Times New Roman', '宋体', 10)

# ============ 英文前置 ============
para(TITLE_EN, 'Times New Roman', '宋体', 22, bold=True, before=12)
para('Author Name1, Author Name2', 'Times New Roman', '宋体', 11, bold=True)
para('(Dept. name of organization, City, Postal Code, Country)', 'Times New Roman', '宋体', 10)
para('Received: ******', 'Times New Roman', '宋体', 10)

ABSTRACT_EN = ('Small and medium-sized enterprises (SMEs) have long relied on scattered local Excel spreadsheets and paper notes '
    'for financial and operational data management, which suffers from formula fragility, cross-sheet duplicate entry, and the inability '
    'to aggregate across stores or periods. This paper presents the design and implementation of a multi-business intelligent management '
    'system based on the Univer online spreadsheet and the Doubao large language model (LLM). The system adopts an "entry-presentation '
    'separation" architecture with PostgreSQL as the single source of truth and Univer as the entry surface only. Its core contribution is '
    'an AI semantic extraction pipeline that serves as the sole data ingestion path: the LLM aligns fields by header text semantics rather '
    'than cell coordinates, replacing fragile positional synchronization with layout-resistant, cross-business-reusable, and near-deterministic '
    'ingestion. A worksheet-level pessimistic lock replaces Operational Transformation for lightweight collaborative entry. The complete system '
    'has been deployed and is in use at real stores. Validation with real footbath-store operational data for July 2026 shows 16 valid business '
    'days successfully ingested, while formula-error cells were correctly rejected by the validation layer.')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
setfont(p.add_run('Abstract: '), 'Times New Roman', '宋体', 12, bold=True)
setfont(p.add_run(ABSTRACT_EN), 'Times New Roman', '宋体', 10)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(12); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
setfont(p.add_run('Keywords: '), 'Times New Roman', '宋体', 12, bold=True)
setfont(p.add_run('small and medium-sized enterprises; business data; large language model; information extraction; online spreadsheet; pessimistic lock'),
        'Times New Roman', '宋体', 10)

# 版权声明
para('Copyright © 2026 by author(s) and Hans Publishers Inc.', 'Times New Roman', '宋体', 9, before=6)
para('This work is licensed under the Creative Commons Attribution International License (CC BY 4.0).', 'Times New Roman', '宋体', 9)
para('http://creativecommons.org/licenses/by/4.0/', 'Times New Roman', '宋体', 9)

# ============ 正文解析 ============
with open('docs/论文-正文草稿.md', encoding='utf-8') as f:
    lines = f.readlines()


def heading_text(text):
    """'1 引言' -> '1. 引言'; '3.1 设计' -> '3.1. 设计'"""
    m = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', text)
    return '%s. %s' % (m.group(1), m.group(2)) if m else text


def add_heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6 if level < 3 else 0)
    r = p.add_run(heading_text(text))
    sz = {1: 12, 2: 11, 3: 10.5}[level]
    setfont(r, 'Times New Roman', '黑体', sz, bold=True)
    return p


def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    parts = re.split(r'(\*\*.+?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); setfont(r, 'Times New Roman', '宋体', 10.5, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); setfont(r, 'Courier New', '宋体', 9)
        else:
            r = p.add_run(part); setfont(r, 'Times New Roman', '宋体', 10.5)
    return p


def add_reference(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.hanging_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text); setfont(r, 'Times New Roman', '宋体', 9)
    return p


FIGS = {
    1: ('fig1-architecture.png', 'Figure 1. System Overall Architecture', '图1. 系统总体架构图'),
    2: ('fig2-footbath-eval.png', 'Figure 2. Real Business Data of Footbath Store (Jul 2026)', '图2. 足浴门店2026-07真实经营数据'),
    3: ('fig3-er-diagram.png', 'Figure 3. Data Model ER Diagram', '图3. 数据模型ER图'),
    4: ('fig4-extraction-pipeline.png', 'Figure 4. AI Extraction Pipeline Flow', '图4. AI抽取管线流程图'),
    5: ('fig5-lock-state-machine.png', 'Figure 5. Pessimistic Lock State Machine', '图5. 悲观锁状态转换图'),
    6: ('fig6-template-five-uses.png', 'Figure 6. Template Descriptor: Five Reuses', '图6. 模板描述符"一份五用"'),
    7: ('fig7-business-dispatch.png', 'Figure 7. Multi-Business Dispatch', '图7. 多业务分派示意图'),
}
_inserted = set()

def insert_figure(n):
    if n in _inserted or n not in FIGS:
        return
    _inserted.add(n)
    fname, cap_en, cap_cn = FIGS[n]
    path = 'docs/%s' % fname
    if not os.path.exists(path):
        return
    para('', 'Times New Roman', '宋体', 10)                                   # 图上方空行
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.add_run().add_picture(path, width=Inches(5.0))
    para(cap_en, 'Times New Roman', '宋体', 9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=3)
    para(cap_cn, 'Times New Roman', '黑体', 9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para('', 'Times New Roman', '宋体', 10)                                   # 图题下空行


i = 0
in_fig_section = False
saw_refs = False
while i < len(lines):
    line = lines[i].rstrip('\n')
    i += 1
    if not line.strip():
        continue
    if line.startswith('# ') or line.startswith('> ') or line.startswith('---'):
        continue
    if line.startswith('**摘要**') or line.startswith('**Abstract**') or line.startswith('**Keywords**'):
        continue
    if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
        continue
    if line.startswith('## 图题') or line.startswith('**图题'):
        in_fig_section = True
        continue
    if in_fig_section and line.startswith('图 '):
        continue
    m_eq = re.match(r'^\$\$(.+)\$\$$', line)
    if m_eq:
        add_equation(doc, m_eq.group(1).strip())
        continue
    if line.startswith('#### '):
        add_heading(line.replace('#### ', ''), 3); continue
    if line.startswith('### '):
        add_heading(line.replace('### ', ''), 2); continue
    if line.startswith('## '):
        in_fig_section = False
        htext = line.replace('## ', '')
        # 参考文献前插入 致谢 / 基金项目 占位
        if htext.strip() == '参考文献' and not saw_refs:
            saw_refs = True
            add_heading('致谢', 1)
            para('（致谢信息。）', 'Times New Roman', '宋体', 10.5, indent=0.74)
            add_heading('基金项目', 1)
            para('（若为基金资助项目,请注明基金名称与编号。）', 'Times New Roman', '宋体', 10.5, indent=0.74)
        add_heading(htext, 1)
        continue
    if re.match(r'^\[\d+\]', line):
        add_reference(line); continue
    if line.startswith('图 '):
        continue
    add_body(line)
    for mfig in re.finditer(r'图\s*(\d+)', line):
        insert_figure(int(mfig.group(1)))

# 未提及的图文末补插
for _n in sorted(FIGS):
    insert_figure(_n)

# ============ 页面 ============
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(28.5)
section.top_margin = Cm(3.0)
section.bottom_margin = Cm(3.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.header_distance = Cm(1.5)
section.footer_distance = Cm(1.5)

out = 'docs/论文-计算机科学与应用版.docx'
doc.save(out)
print('Saved:', out)
