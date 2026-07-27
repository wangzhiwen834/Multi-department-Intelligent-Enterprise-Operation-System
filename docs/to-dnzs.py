# 论文 -> 《电脑知识与技术》稿件模板排版 (python-docx + pandoc 原生 OMML 公式)
# 用法: python docs/to-dnzs.py   (从仓库根目录运行)
# 依据: docs/《电脑知识与技术》稿件模板及要求.doc
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import re
import os
import subprocess
import tempfile
import zipfile
import copy
import shutil

# ============ pandoc LaTeX -> 原生 OMML ============
def _find_pandoc():
    p = shutil.which('pandoc')
    if p:
        return p
    for cand in (r"C:\Users\11422\AppData\Local\Pandoc\pandoc.exe",
                 r"C:\Program Files\Pandoc\pandoc.exe"):
        if os.path.exists(cand):
            return cand
    return None

PANDOC = _find_pandoc()
_M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

def _latex_to_omath(latex):
    if not PANDOC:
        return None
    md = '$' + latex + '$'
    with tempfile.TemporaryDirectory() as td:
        mdp = os.path.join(td, 'eq.md'); docxp = os.path.join(td, 'eq.docx')
        with open(mdp, 'w', encoding='utf-8') as f:
            f.write(md)
        r = subprocess.run([PANDOC, mdp, '-f', 'markdown', '-t', 'docx', '-o', docxp],
                           capture_output=True)
        if r.returncode != 0 or not os.path.exists(docxp):
            return None
        with zipfile.ZipFile(docxp) as z:
            xml = z.read('word/document.xml')
    tree = etree.fromstring(xml)
    return tree.find('.//{%s}oMath' % _M_NS)

_EQ_N = [0]
def add_equation(doc, latex):
    _EQ_N[0] += 1
    p = doc.add_paragraph(style='Normal')
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0); pf.left_indent = Cm(0)
    pf.space_before = Pt(3); pf.space_after = Pt(3)
    pf.tab_stops.add_tab_stop(Cm(7.33), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Cm(14.66), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(); r1.add_tab()
    omath = _latex_to_omath(latex)
    if omath is not None:
        p._element.append(copy.deepcopy(omath))
    else:
        rb = p.add_run(latex); rb.font.name = 'Cambria Math'
        rb._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); rb.font.size = Pt(10.5)
    r2 = p.add_run(); r2.add_tab()
    r3 = p.add_run('(%d)' % _EQ_N[0]); r3.font.size = Pt(10.5)
    return p

# ============ 文档与样式 ============
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
style.paragraph_format.line_spacing = Pt(20)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)


def setfont(run, ascii_font, ea_font, size, bold=None, italic=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), ea_font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def para_centered(text, ea_font, size, bold=False, ascii_font='Times New Roman'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    r = p.add_run(text)
    setfont(r, ascii_font, ea_font, size, bold)
    return p


def label_para(label, body, label_ea='黑体', body_ea='宋体', size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    r1 = p.add_run(label)
    setfont(r1, 'Times New Roman', label_ea, size, bold=True)
    r2 = p.add_run(body)
    setfont(r2, 'Times New Roman', body_ea, size)
    return p


# ============ 前置元数据(模板必填,占位符) ============
# 基金项目 / 作者简介(小五 9pt, 顶格)
p = doc.add_paragraph(); p.paragraph_format.line_spacing = Pt(16)
setfont(p.add_run('基金项目：'), 'Times New Roman', '宋体', 9, bold=True)
p = doc.add_paragraph(); p.paragraph_format.line_spacing = Pt(16)
setfont(p.add_run('作者简介：×××（19××—），男／女，×省×市人，职称，学位，主要研究方向为智能经营系统、大语言模型应用。'),
        'Times New Roman', '宋体', 9, bold=True)

# ============ 题名 / 作者 / 单位 ============
para_centered('多部门智能经营系统的设计与实现', '黑体', 16, bold=True)
para_centered('作者姓名1，作者姓名2', '楷体_GB2312', 10.5)
para_centered('（1. 单位一，省 市 邮编；2. 单位二，省 市 邮编）', '楷体_GB2312', 10.5)

# ============ 摘要 / 关键词 / 中图分类号 ============
ABSTRACT_CN = ('中小企业经营数据长期依赖散乱 Excel 与纸质笔记,存在跨表公式脆弱与重复录入的突出问题。'
    '为此设计并实现了一个基于 Univer 在线表格与豆包大模型的多业态智能经营系统。系统采用"录入与展示分离"架构,'
    '以 PostgreSQL 为唯一事实源、Univer 为录入面;以 AI 语义抽取作为唯一入库路径,以大模型按表头文字语义对齐字段替代位置式同步,'
    '实现抗布局变化、跨业务可复用、近确定性可复现的经营数据入库;并以工作表级悲观锁替代操作变换实现轻量协同录入。'
    '系统已部署上线并交付真实门店使用,以足浴门店真实经营数据验证了抽取管线的有效性,遗留 Excel 的公式错误值被校验机制正确拦截。'
    '研究表明大语言模型可作为结构化数据入库的语义对齐层,为中小企业经营数字化提供了可复用的设计模式。')
label_para('摘要：', ABSTRACT_CN)
label_para('关键词：', '中小企业；经营数据；大语言模型；信息抽取；在线表格；悲观锁')

p = doc.add_paragraph(); p.paragraph_format.line_spacing = Pt(20)
setfont(p.add_run('中图分类号：'), 'Times New Roman', '黑体', 10.5, bold=True)
setfont(p.add_run('TP311    '), 'Times New Roman', '宋体', 10.5)
setfont(p.add_run('文献标识码：'), 'Times New Roman', '黑体', 10.5, bold=True)
setfont(p.add_run('A'), 'Times New Roman', '宋体', 10.5)

# 英文摘要(保留)
ABSTRACT_EN = ('Small and medium-sized enterprises (SMEs) have long relied on scattered local Excel spreadsheets and paper notes '
    'for financial and operational data management, which suffers from formula fragility, cross-sheet duplicate entry, and the inability '
    'to aggregate across stores or periods. This paper presents the design and implementation of a multi-business intelligent management '
    'system based on the Univer online spreadsheet and the Doubao large language model (LLM). The system adopts an "entry-presentation '
    'separation" architecture with PostgreSQL as the single source of truth and Univer as the entry surface only. Its core contribution is '
    'an AI semantic extraction pipeline that serves as the sole data ingestion path: the LLM aligns fields by header text semantics rather '
    'than cell coordinates, replacing fragile positional synchronization with layout-resistant, cross-business-reusable, and near-deterministic '
    'ingestion. A worksheet-level pessimistic lock replaces Operational Transformation for lightweight collaborative entry. The complete system '
    'has been deployed and is in use at real stores. Validation with real footbath-store operational data for July 2026 shows 16 valid business '
    'days successfully ingested, while formula-error cells were correctly rejected by the validation layer.')
label_para('Abstract: ', ABSTRACT_EN, label_ea='Times New Roman', body_ea='Times New Roman')
label_para('Keywords: ', 'small and medium-sized enterprises; business data; large language model; information extraction; online spreadsheet; pessimistic lock',
           label_ea='Times New Roman', body_ea='Times New Roman')

# ============ 正文解析 ============
with open('docs/论文-正文草稿.md', encoding='utf-8') as f:
    lines = f.readlines()


def add_heading(text, level):
    # level 1/2/3 -> 黑体 12pt / 黑体 10.5pt / 宋体 10.5pt, 顶格
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    if level == 1:
        setfont(r, 'Times New Roman', '黑体', 12, bold=True)
    elif level == 2:
        setfont(r, 'Times New Roman', '黑体', 10.5, bold=True)
    else:
        setfont(r, 'Times New Roman', '宋体', 10.5, bold=False)
    return p


def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    parts = re.split(r'(\*\*.+?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); setfont(r, 'Times New Roman', '宋体', 10.5, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); setfont(r, 'Courier New', '宋体', 9)
        else:
            r = p.add_run(part); setfont(r, 'Times New Roman', '宋体', 10.5)
    return p


def add_reference(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.hanging_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(text); setfont(r, 'Times New Roman', '宋体', 9)
    return p


FIGS = {
    1: ('fig1-architecture.png', '图 1  系统总体架构图'),
    2: ('fig2-footbath-eval.png', '图 2  足浴门店 2026-07 真实经营数据'),
    3: ('fig3-er-diagram.png', '图 3  数据模型 ER 图'),
    4: ('fig4-extraction-pipeline.png', '图 4  AI 抽取管线流程图'),
    5: ('fig5-lock-state-machine.png', '图 5  悲观锁状态转换图'),
    6: ('fig6-template-five-uses.png', '图 6  模板描述符"一份五用"'),
    7: ('fig7-business-dispatch.png', '图 7  多业务分派示意图'),
}
_inserted = set()

def insert_figure(n):
    if n in _inserted or n not in FIGS:
        return
    _inserted.add(n)
    fname, caption = FIGS[n]
    path = 'docs/%s' % fname
    if not os.path.exists(path):
        return
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(5.2))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    setfont(cap.add_run(caption), 'Times New Roman', '黑体', 9)

i = 0
in_fig_section = False
while i < len(lines):
    line = lines[i].rstrip('\n')
    i += 1
    if not line.strip():
        continue
    # 跳过前置元数据区(题名/摘要/Abstract/Keywords/草稿说明)与分隔线
    if line.startswith('# '):
        continue
    if line.startswith('> '):
        continue
    if line.startswith('---'):
        continue
    if line.startswith('**摘要**') or line.startswith('**Abstract**') or line.startswith('**Keywords**'):
        continue
    if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
        continue
    # 图题小节与图题行 -> 跳过(图由末尾统一插入)
    if line.startswith('## 图题') or line.startswith('**图题'):
        in_fig_section = True
        continue
    if in_fig_section and (line.startswith('图 ') or line.startswith('图\t')):
        continue
    # 显示公式
    m_eq = re.match(r'^\$\$(.+)\$\$$', line)
    if m_eq:
        add_equation(doc, m_eq.group(1).strip())
        continue
    # 标题
    if line.startswith('#### '):
        add_heading(line.replace('#### ', ''), 3); continue
    if line.startswith('### '):
        add_heading(line.replace('### ', ''), 2); continue
    if line.startswith('## '):
        in_fig_section = False
        add_heading(line.replace('## ', ''), 1); continue
    # 参考文献
    if re.match(r'^\[\d+\]', line):
        add_reference(line); continue
    # 图题行(正文区偶现) -> 跳过
    if line.startswith('图 '):
        continue
    # 普通正文
    add_body(line)
    # 正文提及"图 N"处就地插图(模板要求图表排在正文相应位置)
    for mfig in re.finditer(r'图\s*(\d+)', line):
        insert_figure(int(mfig.group(1)))

# ============ 未在正文提及的图,文末补插 ============
for _n in sorted(FIGS):
    insert_figure(_n)

# ============ 文末联系信息(模板要求) ============
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
setfont(p.add_run('★ 请在文后提供以下信息,方便编辑部后期联系以及邮寄刊物等。'), 'Times New Roman', '宋体', 9, bold=True)
p = doc.add_paragraph()
setfont(p.add_run('1. 联系人;2. 通讯地址(邮政编码);3. 电子信箱、电话。'), 'Times New Roman', '宋体', 9)

# ============ 页面 ============
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

out = 'docs/论文-电脑知识与技术版.docx'
doc.save(out)
print('Saved:', out)
